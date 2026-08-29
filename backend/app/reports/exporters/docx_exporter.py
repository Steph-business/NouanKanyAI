"""
app/reports/exporters/docx_exporter.py — Exportateur de rapports énergétiques au format DOCX (Word).
"""

import io
import logging
from typing import Dict, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.reports.exporters.base import BaseReportExporter
from app.reports.models import EnergyReportData

logger = logging.getLogger("nouankany.reports")


def set_cell_background(cell, hex_color: str):
    """Applique une couleur d'arrière-plan à une cellule de tableau DOCX."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.replace("#", ""))
    tc_pr.append(shd)


class DOCXReportExporter(BaseReportExporter):
    """
    Générateur de documents Microsoft Word (.docx) professionnels.
    """

    def export(
        self,
        report_data: EnergyReportData,
        chart_images: Optional[Dict[str, bytes]] = None,
    ) -> bytes:
        """
        Génère le document DOCX complet.
        """
        doc = Document()

        # Marges standard
        for section in doc.sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

        # 1. En-tête et Titre
        brand_p = doc.add_paragraph()
        brand_run = brand_p.add_run("NOUANKANY.AI — Plateforme d'Intelligence Énergétique Industrielle")
        brand_run.font.name = "Arial"
        brand_run.font.size = Pt(9)
        brand_run.font.color.rgb = RGBColor(2, 132, 199)
        brand_run.font.bold = True

        title_p = doc.add_paragraph()
        title_run = title_p.add_run(report_data.title)
        title_run.font.name = "Arial"
        title_run.font.size = Pt(18)
        title_run.font.color.rgb = RGBColor(15, 23, 42)
        title_run.font.bold = True

        sub_p = doc.add_paragraph()
        sub_text = f"Site : {report_data.site_name} | Période : {report_data.period_start} au {report_data.period_end} | Émis le : {report_data.generated_at}"
        sub_run = sub_p.add_run(sub_text)
        sub_run.font.name = "Arial"
        sub_run.font.size = Pt(9.5)
        sub_run.font.color.rgb = RGBColor(100, 116, 139)

        # 2. Résumé Exécutif
        h1 = doc.add_heading("1. Résumé Exécutif", level=1)
        h1.runs[0].font.color.rgb = RGBColor(2, 132, 199)
        exec_p = doc.add_paragraph(report_data.executive_summary)
        exec_p.style.font.name = "Arial"
        exec_p.style.font.size = Pt(10)

        # 3. Tableau des KPIs
        h2 = doc.add_heading("2. Indicateurs Clés de Performance (KPIs)", level=1)
        h2.runs[0].font.color.rgb = RGBColor(2, 132, 199)
        kpi_table = doc.add_table(rows=4, cols=4)
        kpi_table.style = "Table Grid"

        kpis = report_data.kpis
        data_rows = [
            ("Consommation Totale", f"{kpis.total_energy_kwh:,.1f} kWh", "Facture Globale", f"{kpis.total_cost_fcfa:,.0f} FCFA"),
            ("Consommation Pointe", f"{kpis.peak_hours_energy_kwh:,.1f} kWh", "Coût Pointe (CIE)", f"{kpis.peak_hours_cost_fcfa:,.0f} FCFA"),
            ("Économies Effacement", f"+{kpis.peak_shaving_savings_fcfa:,.0f} FCFA", "Facteur de Puissance", f"cos φ = {kpis.average_power_factor_cos_phi:.2f}"),
            ("Puissance Max", f"{kpis.max_peak_power_kw:.1f} kW", "Anomalies", f"{kpis.anomaly_incidents_count} incident(s)"),
        ]

        for r_idx, row in enumerate(data_rows):
            for c_idx, val in enumerate(row):
                cell = kpi_table.cell(r_idx, c_idx)
                cell.text = val
                cell.paragraphs[0].runs[0].font.name = "Arial"
                cell.paragraphs[0].runs[0].font.size = Pt(9)
                if c_idx % 2 == 0:
                    set_cell_background(cell, "F1F5F9")
                    cell.paragraphs[0].runs[0].font.bold = True
                else:
                    set_cell_background(cell, "FFFFFF")

        doc.add_paragraph("")

        # 4. Graphiques
        if chart_images:
            h3 = doc.add_heading("3. Analyse Graphique", level=1)
            h3.runs[0].font.color.rgb = RGBColor(2, 132, 199)

            if "load_curve" in chart_images:
                doc.add_picture(io.BytesIO(chart_images["load_curve"]), width=Inches(6.5))
                doc.add_paragraph("Figure 1 : Courbe de charge horaire et puissance souscrite.")

            if "tariff_pie" in chart_images:
                doc.add_picture(io.BytesIO(chart_images["tariff_pie"]), width=Inches(5.0))
                doc.add_paragraph("Figure 2 : Répartition par tranche tarifaire CIE.")

        # 5. Détail des Machines
        if report_data.machines:
            h4 = doc.add_heading("4. Bilan par Équipement", level=1)
            h4.runs[0].font.color.rgb = RGBColor(2, 132, 199)
            mach_table = doc.add_table(rows=1 + len(report_data.machines[:8]), cols=6)
            mach_table.style = "Table Grid"

            headers = ["Machine", "Catégorie", "Énergie (kWh)", "Coût (FCFA)", "Heures", "Note"]
            for c_idx, h in enumerate(headers):
                cell = mach_table.cell(0, c_idx)
                cell.text = h
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                set_cell_background(cell, "0F172A")

            for r_idx, m in enumerate(report_data.machines[:8], start=1):
                row_vals = [m.machine_name, m.category, f"{m.energy_kwh:,.1f}", f"{m.cost_fcfa:,.0f}", f"{m.running_hours:.1f}h", m.efficiency_grade]
                for c_idx, val in enumerate(row_vals):
                    cell = mach_table.cell(r_idx, c_idx)
                    cell.text = val
                    cell.paragraphs[0].runs[0].font.size = Pt(8.5)
                    if r_idx % 2 == 0:
                        set_cell_background(cell, "F8FAFC")

            doc.add_paragraph("")

        # 6. Recommandations IA
        if report_data.recommendations:
            h5 = doc.add_heading("5. Plan d'Action & Recommandations IA", level=1)
            h5.runs[0].font.color.rgb = RGBColor(2, 132, 199)
            for rec in report_data.recommendations:
                p = doc.add_paragraph(style="List Bullet")
                r_title = p.add_run(f"[{rec.title}] (Gains : +{rec.estimated_savings_fcfa:,.0f} FCFA / {rec.estimated_savings_kwh:,.0f} kWh)\n")
                r_title.bold = True
                r_desc = p.add_run(f"{rec.description} (Cible : {rec.target_equipment})")
                r_desc.font.size = Pt(9.5)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()
